"""OCR + 结构化 Pipeline（多模型支持）。

主流程 process_note(note_id, model_id=None):
  1. 加载笔记行
  2. 文件转图像（PDF 用 PyMuPDF，PNG/JPG 直接读字节）
  3. 调多模态 vision 模型抽取 {title, ocr_text, summary, keywords[]}
     - 若指定 model_id，仅用该模型
     - 否则用 primary 模型，失败时按 enabled 顺序 fallback
  4. 调 embedding 模型生成向量
  5. 写库（含 ocr_model 字段）+ 生成缩略图
  6. 触发 graph_api.recompute_links_for_note

当未配置 OPENAI_API_KEY 时进入 DEMO 模式：
  - title 取文件名
  - ocr_text 写入提示语
  - keywords 为空
  - embedding 为随机 1536 维向量（保证图谱/检索端到端可用）
"""
from __future__ import annotations

import base64
import hashlib
import json
import logging
import os
import random
from typing import Any, Dict, List, Optional

import database
import graph_api
import settings_store
from config import get_config

logger = logging.getLogger("brain.ocr")

SUPPORTED_EXTS = (".pdf", ".png", ".jpg", ".jpeg", ".txt", ".md", ".markdown", ".docx")
# 文本型扩展名（不经过 OCR vision，直接抽文本 → LLM 结构化）
TEXT_EXTS = (".txt", ".md", ".markdown", ".docx")


# ---------------------------------------------------------------------------
# 文件工具
# ---------------------------------------------------------------------------
def compute_file_hash(path: str, chunk_size: int = 1 << 16) -> str:
    """计算文件 SHA256 哈希，用于去重。"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def file_to_images(path: str) -> List[str]:
    """把笔记文件转成 base64 编码的图像列表。

    - PDF：每页渲染为 PNG（PyMuPDF/fitz）
    - PNG/JPG：直接读取字节并 base64
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
        except ImportError as e:  # pragma: no cover
            raise RuntimeError("PyMuPDF(fitz) 未安装，无法处理 PDF") from e
        images: List[str] = []
        doc = fitz.open(path)
        try:
            # 控制总页数，避免超大 PDF 把 token 打爆
            max_pages = min(doc.page_count, 10)
            for i in range(max_pages):
                page = doc.load_page(i)
                # 150 DPI 左右，对 OCR 足够
                zoom = 2.0
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                png_bytes = pix.tobytes("png")
                images.append(base64.b64encode(png_bytes).decode("ascii"))
        finally:
            doc.close()
        return images
    elif ext in (".png", ".jpg", ".jpeg"):
        with open(path, "rb") as f:
            data = f.read()
        return [base64.b64encode(data).decode("ascii")]
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def extract_text_from_file(path: str) -> str:
    """从文本型文件中提取纯文本。

    支持：
      - .txt / .md / .markdown：直接读取（自动尝试 UTF-8 / GBK）
      - .docx：用 python-docx 提取段落 + 表格文本

    返回纯文本字符串；失败抛异常。
    """
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md", ".markdown"):
        # 优先 UTF-8，回退 GBK（Windows 记事本常见）
        for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise RuntimeError("无法解码文本文件（尝试 UTF-8/GBK 均失败）")
    if ext == ".docx":
        try:
            from docx import Document
        except ImportError as e:  # pragma: no cover
            raise RuntimeError("python-docx 未安装，无法处理 .docx") from e
        doc = Document(path)
        parts: List[str] = []
        # 段落
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    raise ValueError(f"不支持的文本型文件: {ext}")


def generate_thumbnail(path: str, out_path: str, width: int = 200, quality: int = 80) -> str:
    """生成缩略图（200px 宽，JPEG 质量 80）。返回输出路径。

    - PDF：第一页渲染
    - 图片：直接缩放
    - 文本型（txt/md/docx）：用 PIL 把前若干行文本绘制到白色画布上
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("Pillow 未安装，无法生成缩略图") from e

    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in TEXT_EXTS:
            # 文本型缩略图：白底黑字，绘制前 ~24 行
            try:
                text = extract_text_from_file(path)
            except Exception as e:
                logger.warning("文本缩略图抽取失败 %s: %s", path, e)
                return ""
            lines = [ln for ln in text.splitlines() if ln.strip()][:24]
            if not lines:
                lines = ["(空文件)"]
            # 画布尺寸：宽 400px（再缩放到 width），行高 18px
            canvas_w = 400
            line_h = 20
            canvas_h = max(120, line_h * (len(lines) + 2))
            img = Image.new("RGB", (canvas_w, canvas_h), color=(255, 255, 255))
            draw = ImageDraw.Draw(img)
            # 使用默认字体（部署环境通常无中文字体，退化到 bitmap）
            try:
                font = ImageFont.truetype("DejaVuSans.ttf", 14)
            except Exception:
                font = ImageFont.load_default()
            y = line_h
            for ln in lines:
                # 截断过长行
                draw.text((10, y), ln[:50], fill=(30, 30, 30), font=font)
                y += line_h
            # 缩放到目标宽度
            ratio = width / float(canvas_w)
            new_h = max(1, int(canvas_h * ratio))
            thumb = img.resize((width, new_h), Image.LANCZOS)
            os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
            thumb.save(out_path, "JPEG", quality=quality)
            return out_path

        if ext == ".pdf":
            import fitz
            doc = fitz.open(path)
            try:
                if doc.page_count == 0:
                    return ""
                page = doc.load_page(0)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                png_bytes = pix.tobytes("png")
            finally:
                doc.close()
            import io
            img = Image.open(io.BytesIO(png_bytes)).convert("RGB")
        else:
            img = Image.open(path).convert("RGB")

        ratio = width / float(img.width) if img.width else 1.0
        new_height = max(1, int(img.height * ratio))
        thumb = img.resize((width, new_height), Image.LANCZOS)
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
        thumb.save(out_path, "JPEG", quality=quality)
        return out_path
    except Exception as e:
        logger.warning("缩略图生成失败 %s: %s", path, e)
        return ""


# ---------------------------------------------------------------------------
# OpenAI 客户端
# ---------------------------------------------------------------------------
def _get_client():
    """惰性构造 OpenAI 客户端。未配置 key 时返回 None（进入 demo 模式）。"""
    cfg = get_config()
    if not cfg.OPENAI_API_KEY:
        return None
    try:
        from openai import OpenAI
    except ImportError as e:  # pragma: no cover
        logger.error("openai SDK 未安装: %s", e)
        return None
    kwargs: Dict[str, Any] = {"api_key": cfg.OPENAI_API_KEY}
    if cfg.OPENAI_BASE_URL:
        kwargs["base_url"] = cfg.OPENAI_BASE_URL
    return OpenAI(**kwargs)


def _strip_fences(text: str) -> str:
    """剥离 ```json ... ``` 等 markdown 代码围栏。"""
    s = text.strip()
    if s.startswith("```"):
        # 去掉首行 ```json / ```
        first_nl = s.find("\n")
        if first_nl != -1:
            s = s[first_nl + 1:]
        if s.endswith("```"):
            s = s[: -3]
    return s.strip()


def _parse_structured(raw: str) -> Dict[str, Any]:
    """健壮地解析 LLM 返回的结构化 JSON。"""
    cleaned = _strip_fences(raw)
    # 尝试直接解析
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # 尝试抽取第一个 {...}
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(cleaned[start: end + 1])
        except json.JSONDecodeError:
            pass
    return {}


# ---------------------------------------------------------------------------
# OCR + 结构化（真实 / Demo）
# ---------------------------------------------------------------------------
_OCR_PROMPT = """你是一名专业的手写笔记 OCR 与结构化助手。请仔细识别图片中的所有手写内容，要求：

1. **完整准确地转录**所有手写文字，包括公式、符号、图表标注等。
2. 保留原文的段落结构与层次（标题、列表、缩进）。
3. 数学公式用 LaTeX 语法（行内 $...$，独立块 $$...$$）。
4. 表格用 Markdown 表格语法。
5. 无法识别的字用 □ 占位，不要瞎猜。
6. 若有多张图，按页顺序输出，页与页之间用 `---PAGE---` 分隔。

**手写元素语义标注**（重要）：
- **箭头指向**：手写箭头连接两个概念时，用 `[→: 起点指向终点]` 标注。
  例：知识 → 实践 写成 `[→: 知识指向实践]`
- **划线删除**：被划掉（删除线）的内容用 markdown 删除线：`~~删除的字~~`。
  保留原字以便追溯，但用删除线标记为已删除。
- **小字批注**：行间或旁边的批注、补充说明，用 `[批注: 批注内容]` 标注。
  批注应就近放在被批注的段落后面。
- **下划线/重点标记**：用下划线的重点字词用 `[重点: 内容]` 标注。
- **圈选/方框**：被圈起来或方框包围的关键字用 `[圈选: 内容]` 标注。

识别完成后，再用同样的 JSON 格式返回结构化字段：

```json
{
  "title": "简短标题（5-15字，概括主题）",
  "ocr_text": "完整识别的文本（含公式/表格/语义标签的 markdown）",
  "summary": "1-3 句摘要，提炼核心知识点",
  "keywords": ["关键词1", "关键词2", ...]
}
```

要求：
- 仅返回上述 JSON 对象，不要任何额外解释或前后缀
- 关键词 3-8 个，涵盖主题、方法、对象等
- title 不要包含"笔记""note"等无意义词
- ocr_text 必须保留全部识别内容，不要截断"""


def _call_vision_model(client, model_id: str, images: List[str]) -> Dict[str, Any]:
    """调用指定的多模态模型做 OCR + 结构化。

    单图（图片笔记或单页 PDF）：一次调用，结构化输出。
    多图（多页 PDF）：分页调用，每页单独 OCR，最后合并 ocr_text，
    title 取首页，summary 综合各页摘要。这样能避免多图一次性发送导致
    模型注意力分散、内容丢失或被 token 上限截断。

    Args:
        client: OpenAI 客户端
        model_id: 模型 ID（如 Qwen/Qwen3-VL-32B-Instruct）
        images: base64 编码的图片列表

    Returns:
        解析后的结构化 dict（可能为空 dict 表示解析失败）
    """
    if len(images) <= 1:
        return _call_vision_single(client, model_id, images)

    # 多页：分页 OCR，最后合并
    logger.info("多页文档（%d 页），分页 OCR 后合并", len(images))
    page_results: List[Dict[str, Any]] = []
    for idx, img in enumerate(images):
        try:
            r = _call_vision_single(client, model_id, [img], page_idx=idx + 1, total_pages=len(images))
            if r:
                page_results.append(r)
            else:
                logger.warning("第 %d/%d 页 OCR 返回空，跳过", idx + 1, len(images))
        except Exception as e:
            logger.warning("第 %d/%d 页 OCR 失败: %s", idx + 1, len(images), e)
            continue

    if not page_results:
        return {}

    # 合并
    merged_ocr_parts: List[str] = []
    for i, r in enumerate(page_results, 1):
        page_text = r.get("ocr_text") or ""
        if page_text:
            merged_ocr_parts.append(f"--- 第 {i} 页 ---\n{page_text}")
    merged_ocr = "\n\n".join(merged_ocr_parts)

    # title 用第一页的
    merged_title = page_results[0].get("title") or ""
    # summary 综合各页
    summary_parts = [r.get("summary") for r in page_results if r.get("summary")]
    merged_summary = " | ".join(summary_parts) if summary_parts else ""

    # keywords 合并去重
    kw_set = []
    for r in page_results:
        for k in r.get("keywords") or []:
            if k not in kw_set:
                kw_set.append(k)

    return {
        "title": merged_title or "(多页笔记)",
        "ocr_text": merged_ocr,
        "summary": merged_summary,
        "keywords": kw_set[:10],
    }


def _call_vision_single(client, model_id: str, images: List[str],
                        page_idx: Optional[int] = None, total_pages: Optional[int] = None) -> Dict[str, Any]:
    """单次调用视觉模型（单图或单页）。

    Args:
        client: OpenAI 客户端
        model_id: 模型 ID
        images: 单张 base64 图片
        page_idx: 当前页码（多页时用）
        total_pages: 总页数（多页时用）
    """
    if page_idx and total_pages:
        prompt = f"""你是一名专业的手写笔记 OCR 与结构化助手。这是多页文档的第 {page_idx}/{total_pages} 页，请仔细识别本页的所有手写内容，要求：

1. **完整准确地转录**本页所有手写文字，包括公式、符号、图表标注等。
2. 保留原文的段落结构与层次（标题、列表、缩进）。
3. 数学公式用 LaTeX 语法（行内 $...$，独立块 $$...$$）。
4. 表格用 Markdown 表格语法。
5. 无法识别的字用 □ 占位，不要瞎猜。

**手写元素语义标注**（重要）：
- **箭头指向**：手写箭头连接两个概念时，用 `[→: 起点指向终点]` 标注。
- **划线删除**：被划掉的内容用 `~~删除的字~~` 标注，保留原字。
- **小字批注**：行间或旁边的批注用 `[批注: 批注内容]` 标注，就近放。
- **下划线重点**：用 `[重点: 内容]` 标注。
- **圈选方框**：用 `[圈选: 内容]` 标注。

识别完成后，用 JSON 返回本页的结构化字段：

```json
{{
  "title": "本页标题（5-15字）",
  "ocr_text": "本页完整识别的文本（含语义标签）",
  "summary": "本页 1-2 句摘要",
  "keywords": ["本页关键词1", "本页关键词2"]
}}
```

要求：
- 仅返回上述 JSON 对象，不要任何额外解释
- ocr_text 必须保留本页全部识别内容，不要截断"""
    else:
        prompt = _OCR_PROMPT

    content: List[Dict[str, Any]] = [{"type": "text", "text": prompt}]
    for img_b64 in images[:1]:  # 单次只发一张图
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{img_b64}"},
        })
    # 注入用户习惯修正示例（如果有）
    correction_hint = _build_ocr_correction_hint()
    if correction_hint:
        content.insert(0, {"type": "text", "text": correction_hint})
    resp = client.chat.completions.create(
        model=model_id,
        messages=[{"role": "user", "content": content}],
        temperature=0.1,
        max_tokens=8000,  # 单页 OCR 上限放宽
        timeout=120,
    )
    raw = resp.choices[0].message.content or ""
    return _parse_structured(raw)


def _ocr_structured(client, images: List[str], model_id: Optional[str] = None, file_path: Optional[str] = None) -> tuple[Dict[str, Any], Optional[str]]:
    """OCR + 结构化，支持指定模型与 fallback。

    Args:
        client: OpenAI 客户端
        images: base64 图片列表（仅 OpenAI vision 路径用）
        model_id: 指定使用的模型 id（settings_store 中的 id）。
                  None 表示用 primary，失败时 fallback。
        file_path: 原始文件路径（百度 OCR 需要，直接读文件而非 base64）

    Returns:
        (structured_dict, used_model_id)
        - structured_dict 解析失败时为空 dict
        - used_model_id 实际成功调用的模型 id（settings_store 里的 id），
          失败时为 None
    """
    # baidu 模型走专用路径，不调 OpenAI vision
    if model_id == "baidu" or (model_id is None and _should_try_baidu_first()):
        result = _try_baidu_ocr(file_path)
        if result:
            return result, "baidu"
        # baidu 失败：如果指定了 baidu，直接返回失败；
        # 如果是 fallback 路径（model_id=None），继续往下试 OpenAI 模型
        if model_id == "baidu":
            return {}, None

    if model_id and model_id != "baidu":
        # 指定模型：只试这一个
        m = settings_store.get_ocr_model_by_id(model_id)
        if not m:
            logger.warning("指定的 OCR 模型 %s 不存在", model_id)
            return {}, None
        try:
            result = _call_vision_model(client, m["model"], images)
            if result:
                return result, m["id"]
            logger.warning("模型 %s 返回内容无法解析", m["model"])
            return {}, None
        except Exception as e:
            logger.warning("模型 %s (%s) 调用失败: %s", m.get("name"), m["model"], e)
            return {}, None

    # 未指定模型：按 enabled 顺序尝试，primary 在前
    candidates = settings_store.get_enabled_ocr_models()
    # 过滤掉 baidu（已经试过了）
    candidates = [m for m in candidates if m.get("id") != "baidu"]
    if not candidates:
        # 回退到 env 默认
        cfg = get_config()
        try:
            result = _call_vision_model(client, cfg.LLM_MODEL, images)
            return result, None
        except Exception as e:
            logger.warning("默认模型 %s 调用失败: %s", cfg.LLM_MODEL, e)
            return {}, None

    last_err: Optional[Exception] = None
    for m in candidates:
        try:
            logger.info("尝试 OCR 模型: %s (%s)", m.get("name"), m["model"])
            result = _call_vision_model(client, m["model"], images)
            if result:
                logger.info("OCR 成功，使用模型: %s", m["model"])
                return result, m["id"]
            logger.warning("模型 %s 返回内容无法解析，尝试下一个", m["model"])
        except Exception as e:
            last_err = e
            logger.warning("模型 %s (%s) 失败，尝试下一个: %s",
                           m.get("name"), m["model"], e)
            continue
    if last_err:
        logger.error("所有 OCR 模型均失败，最后错误: %s", last_err)
    return {}, None


def _should_try_baidu_first() -> bool:
    """判断是否应该优先用百度 OCR（已配置且 baidu 模型在 enabled 列表）。"""
    cfg = get_config()
    if not (cfg.BAIDU_OCR_ENABLED and cfg.BAIDU_OCR_API_KEY and cfg.BAIDU_OCR_SECRET_KEY):
        return False
    baidu_model = settings_store.get_ocr_model_by_id("baidu")
    return bool(baidu_model and baidu_model.get("enabled"))


def _try_baidu_ocr(file_path: Optional[str]) -> Optional[Dict[str, Any]]:
    """调用百度 OCR 处理文件，返回结构化结果。

    流程：
    1. 百度 OCR 识别手写文字（精度高，但不提供结构化字段）
    2. 用 Kimi 等文本模型对百度返回的纯文本做二次结构化
       （不传图，便宜），生成 title/summary/keywords + 语义标签增强
    3. Kimi 失败时回退到本地规则派生（保证可用性）
    """
    if not file_path or not os.path.exists(file_path):
        logger.warning("百度 OCR：文件路径无效或不存在 %s", file_path)
        return None
    cfg = get_config()
    if not (cfg.BAIDU_OCR_API_KEY and cfg.BAIDU_OCR_SECRET_KEY):
        logger.warning("百度 OCR 未配置凭证")
        return None
    try:
        import baidu_ocr
        # 打印文件诊断信息（便于排查 image format error）
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else -1
        file_ext = os.path.splitext(file_path)[1].lower()
        logger.info("尝试 OCR 模型: 百度手写 OCR (handwriting) | file=%s ext=%s size=%d bytes",
                    os.path.basename(file_path), file_ext, file_size)
        text = baidu_ocr.recognize_file(
            file_path,
            api_key=cfg.BAIDU_OCR_API_KEY,
            secret_key=cfg.BAIDU_OCR_SECRET_KEY,
        )
        if not text or not text.strip():
            logger.warning("百度 OCR 返回空文本")
            return None
        text = text.strip()

        # 二次结构化：让 Kimi 基于纯文本生成 title/summary/keywords + 语义标签
        structured = _refine_baidu_text_with_llm(text)
        if structured:
            # 用 LLM 结果，但 ocr_text 保留 LLM 增强后的版本（带语义标签）
            return structured

        # LLM 失败：回退到本地规则
        logger.warning("百度 OCR 二次结构化失败，回退到本地规则")
        first_line = text.split("\n", 1)[0].strip()[:30]
        title = first_line or "(百度 OCR 笔记)"
        summary = text[:100].replace("\n", " ")
        return {
            "title": title,
            "ocr_text": text,
            "summary": summary,
            "keywords": [],
        }
    except Exception as e:
        logger.warning("百度 OCR 失败: %s", e)
        return None


# 让 LLM 对百度 OCR 返回的纯文本做二次结构化（不传图，便宜）
_REFINE_PROMPT_TEMPLATE = """你是一名手写笔记结构化助手。下面是百度 OCR 识别出的手写笔记纯文本（可能含识别误差）。
请基于这段文本生成结构化字段，要求：

1. **修正明显的 OCR 错误**（如错字、断词），但不要改变原意。
2. **标注手写元素**（如果原文里有迹象）：
   - 被划线删除的内容用 markdown 删除线：`~~删除的字~~`
   - 行间/旁批小字批注用：`[批注: 批注内容]`
   - 箭头指向关系用：`[→: 起点指向终点]`
3. 保留原文段落结构（标题、列表、缩进）。
4. 数学公式用 LaTeX，表格用 Markdown 表格。

返回 JSON：
```json
{{
  "title": "5-15字标题",
  "ocr_text": "完整文本（含上述语义标签）",
  "summary": "1-3 句摘要",
  "keywords": ["关键词1", "关键词2"]
}}
```

仅返回 JSON 对象，不要额外解释。

百度 OCR 文本：
```
{raw_text}
```
"""


def _refine_baidu_text_with_llm(raw_text: str) -> Optional[Dict[str, Any]]:
    """用 Kimi 等文本模型对百度 OCR 纯文本做二次结构化。

    不传图，只传文本，所以便宜快速。
    """
    client = _get_client()
    if client is None:
        return None
    cfg = get_config()
    try:
        # 注入用户习惯修正（如果有的话）
        correction_hint = _build_ocr_correction_hint()
        prompt = _REFINE_PROMPT_TEMPLATE.format(raw_text=raw_text[:6000])
        if correction_hint:
            prompt = correction_hint + "\n\n" + prompt

        resp = client.chat.completions.create(
            model=cfg.LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=4000,
            timeout=60,
        )
        raw = resp.choices[0].message.content or ""
        result = _parse_structured(raw)
        if result and result.get("ocr_text"):
            logger.info("百度 OCR 文本经 LLM 二次结构化成功")
            return result
        logger.warning("LLM 二次结构化返回内容无法解析")
        return None
    except Exception as e:
        logger.warning("LLM 二次结构化失败: %s", e)
        return None


def _build_ocr_correction_hint() -> str:
    """从 user_memory 表读取用户习惯，拼成 prompt 提示。

    同时检索两类记忆：
      - type='ocr_correction'：用户修正过的 OCR 错误（错字、改写）
      - type='ocr_addition'：用户过去补充的内容（批注、思考、扩展）

    让 LLM 在 OCR 时：
      1. 避免类似 OCR 错误
      2. 保留用户习惯的补充内容模式（如喜欢在末尾加思考）
    """
    try:
        parts = []

        # 1. 修正示例（top 5）
        corrections = database.list_memory(type="ocr_correction", limit=5)
        if corrections:
            lines = ["以下是用户过去修正过的 OCR 错误示例，请避免类似错误："]
            for c in corrections:
                content = c.get("content", "")
                weight = c.get("weight", 0.5)
                lines.append(f"- {content}  (权重 {weight:.2f})")
            parts.append("\n".join(lines))

        # 2. 补充内容示例（top 3）
        additions = database.list_memory(type="ocr_addition", limit=3)
        if additions:
            lines = ["以下是用户过去在笔记中补充过的内容示例，请参考用户的笔记习惯（如喜欢补充思考、批注等）："]
            for a in additions:
                content = a.get("content", "")
                weight = a.get("weight", 0.5)
                lines.append(f"- {content}  (权重 {weight:.2f})")
            parts.append("\n".join(lines))

        if not parts:
            return ""
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("读取 OCR 习惯记忆失败: %s", e)
        return ""


def _embed_text(client, text: str) -> List[float]:
    """调用 text-embedding-3-small 生成向量。"""
    cfg = get_config()
    resp = client.embeddings.create(model=cfg.EMBEDDING_MODEL, input=text)
    return list(resp.data[0].embedding)


def _demo_structured(file_path: str) -> Dict[str, Any]:
    """Demo 模式：从文件名派生 title，其余占位。"""
    base = os.path.splitext(os.path.basename(file_path))[0]
    return {
        "title": base.replace("_", " ").replace("-", " ").strip() or "(未命名笔记)",
        "ocr_text": "(demo 模式 — 设置 OPENAI_API_KEY 以启用真实 OCR)",
        "summary": f"演示笔记：{base}",
        "keywords": [],
    }


# 文本型结构化 prompt：用于 txt/md/docx 抽取后的 LLM 结构化
_TEXT_STRUCT_PROMPT = """你是一名笔记结构化助手。下面是用户上传的文本型笔记（TXT/Markdown/Word 抽取后的纯文本）。
请基于这段文本生成结构化字段，要求：

1. **完整保留原文**，不要删减或改写内容（ocr_text 字段必须是原文全文）。
2. 推断一个简洁的标题（5-20 字），如果原文开头有明显的 `#` 标题或首行短句，优先用作标题。
3. 生成 1-2 句话的摘要。
4. 提取 3-8 个关键词。
5. 保留原文的 Markdown 语法（标题、列表、代码块、表格等），不要转换格式。
6. 如果原文里有明显的重点标注（如下划线、加粗、高亮），用 `[重点: 内容]` 标注。

用 JSON 返回：

```json
{{
  "title": "标题",
  "ocr_text": "原文全文",
  "summary": "1-2 句摘要",
  "keywords": ["关键词1", "关键词2"]
}}
```

仅返回 JSON，不要任何额外解释。"""


def _structured_from_text(client, text: str, model_id: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """对纯文本做 LLM 结构化（不传图，便宜）。

    用于 TXT/Markdown/DOCX 等文本型文件：先本地抽文本，再让 LLM 生成
    title/summary/keywords。原文本作为 ocr_text 保留。

    Args:
        client: OpenAI 客户端
        text: 抽取出的纯文本
        model_id: 指定模型 id（None 则用 primary LLM 模型）

    Returns:
        结构化 dict，失败返回 None。
    """
    if not text or not text.strip():
        logger.warning("文本结构化：内容为空")
        return None
    cfg = get_config()
    # 选择模型：优先用指定的，否则用 LLM_MODEL（文本任务用通用模型即可）
    if model_id and model_id != "baidu":
        m = settings_store.get_ocr_model_by_id(model_id)
        model_name = m["model"] if m else cfg.LLM_MODEL
    else:
        model_name = cfg.LLM_MODEL
    # 截断超长文本（避免 token 爆炸）
    max_chars = 24000
    truncated = text[:max_chars]
    if len(text) > max_chars:
        logger.info("文本结构化：原文过长（%d 字），截断到 %d 字", len(text), max_chars)
    # 注入用户习惯
    hint = _build_ocr_correction_hint()
    sys_prompt = (_TEXT_STRUCT_PROMPT + "\n\n" + hint) if hint else _TEXT_STRUCT_PROMPT
    try:
        resp = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": f"以下是笔记原文：\n\n{truncated}"},
            ],
            temperature=0.1,
            max_tokens=8000,
            timeout=120,
        )
        raw = resp.choices[0].message.content or ""
        result = _parse_structured(raw)
        if result and result.get("ocr_text"):
            return result
        logger.warning("文本结构化返回内容无法解析")
        return None
    except Exception as e:
        logger.warning("文本结构化 LLM 调用失败: %s", e)
        return None


def _fallback_text_structured(text: str, file_path: str) -> Dict[str, Any]:
    """LLM 不可用时的本地兜底结构化。

    - title：取第一行非空文本（截断 30 字）
    - summary：取前 100 字
    - keywords：空
    - ocr_text：原文
    """
    base = os.path.splitext(os.path.basename(file_path))[0]
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    first_line = lines[0][:30] if lines else base
    title = first_line.lstrip("#").strip() or base
    summary = text[:100].replace("\n", " ").strip()
    return {
        "title": title,
        "ocr_text": text,
        "summary": summary,
        "keywords": [],
    }


def _demo_embedding(seed: int = 42) -> List[float]:
    """Demo 模式：生成确定性的随机 1536 维向量。

    Args:
        seed: 随机种子，建议传入 note_id 使不同笔记得到不同向量，
              从而在 demo 模式下也能形成有差异的图谱。
    """
    cfg = get_config()
    dim = cfg.EMBEDDING_DIM
    rng = random.Random(seed)
    vec = [rng.uniform(-1.0, 1.0) for _ in range(dim)]
    norm = sum(x * x for x in vec) ** 0.5
    if norm > 0:
        vec = [x / norm for x in vec]
    return vec


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def process_note(note_id: int, model_id: Optional[str] = None) -> bool:
    """处理单条笔记的完整 OCR + 结构化 + embedding + 图谱重算流程。

    Args:
        note_id: 笔记 id
        model_id: 指定使用哪个 OCR 模型（settings_store 中的 id）。
                  None 表示用 primary 模型，失败时 fallback。

    Returns:
        True 表示处理成功，False 表示失败。

    任何异常都把状态置为 'failed' 并记录日志，不向上抛出（用于后台 worker）。
    若笔记被人工编辑过（manually_edited=1），跳过 OCR，只重算 embedding 和链接。
    """
    note = database.get_note(note_id)
    if not note:
        logger.warning("process_note: 笔记 %s 不存在", note_id)
        return False
    file_path = note["file_path"]
    if not os.path.exists(file_path):
        logger.error("文件不存在: %s", file_path)
        database.update_note_status(note_id, "failed")
        return False

    # 人工编辑过的笔记：跳过 OCR，只重算 embedding 和链接
    if note.get("manually_edited"):
        logger.info("笔记 %s 已被人工编辑，跳过 OCR，仅重算 embedding/链接", note_id)
        database.update_note_status(note_id, "processing")
        client = _get_client()
        try:
            if client is None:
                embedding = _demo_embedding(seed=note_id)
            else:
                embed_input = (
                    (note.get("title") or "")
                    + "\n"
                    + (note.get("summary") or "")
                    + "\n"
                    + (note.get("ocr_text") or "")
                )
                embedding = _embed_text(client, embed_input)
                database.update_note_fields(note_id, embedding=embedding)
            # 重算链接
            try:
                graph_api.recompute_links_for_note(note_id)
            except Exception as ge:
                logger.warning("链接重算失败 note %s: %s", note_id, ge)
            database.update_note_status(note_id, "done")
            logger.info("笔记 %s 人工编辑版本已重算 embedding (model=manual)", note_id)
            return True
        except Exception as e:
            logger.exception("笔记 %s 人工编辑版本重算失败: %s", note_id, e)
            database.update_note_status(note_id, "failed")
            return False

    database.update_note_status(note_id, "processing")
    client = _get_client()
    is_demo = client is None

    try:
        # 文本型文件分支：直接抽文本 → LLM 结构化（不走 vision 模型）
        ext = os.path.splitext(file_path)[1].lower()
        if ext in TEXT_EXTS:
            logger.info("处理文本型笔记 %s (%s)", note_id, file_path)
            raw_text = extract_text_from_file(file_path)
            if not raw_text or not raw_text.strip():
                raise RuntimeError("文本文件内容为空")

            used_model_id: Optional[str] = None
            if is_demo:
                structured = _fallback_text_structured(raw_text, file_path)
                embedding = _demo_embedding(seed=note_id)
            else:
                structured = _structured_from_text(client, raw_text, model_id=model_id)
                if structured:
                    # 文本型用 LLM_MODEL，标记为 "text-llm" 便于区分
                    used_model_id = "text-llm"
                else:
                    # LLM 失败：回退到本地结构化
                    logger.warning("文本型 LLM 结构化失败，回退到本地规则")
                    structured = _fallback_text_structured(raw_text, file_path)
                    used_model_id = "text-fallback"
                embed_input = (
                    (structured.get("title") or "")
                    + "\n"
                    + (structured.get("summary") or "")
                    + "\n"
                    + (structured.get("ocr_text") or "")
                )
                embedding = _embed_text(client, embed_input)
        else:
            # 图像型文件分支（PDF/PNG/JPG）：走原有 vision OCR 路径
            images = file_to_images(file_path)
            if not images:
                raise RuntimeError("未能从文件提取到任何图像")

            used_model_id = None
            if is_demo:
                logger.info("[demo] 处理笔记 %s (%s)", note_id, file_path)
                structured = _demo_structured(file_path)
                embedding = _demo_embedding(seed=note_id)
            else:
                structured, used_model_id = _ocr_structured(client, images, model_id=model_id, file_path=file_path)
                if not structured:
                    raise RuntimeError("LLM 返回内容无法解析为 JSON（所有候选模型均失败）")
                embed_input = (
                    (structured.get("title") or "")
                    + "\n"
                    + (structured.get("summary") or "")
                    + "\n"
                    + (structured.get("ocr_text") or "")
                )
                embedding = _embed_text(client, embed_input)

        # 缩略图
        cfg = get_config()
        thumb_name = f"{note_id}.jpg"
        thumb_path = os.path.join(cfg.THUMBNAIL_DIR, thumb_name)
        generate_thumbnail(file_path, thumb_path)

        database.update_note_content(
            note_id,
            title=structured.get("title") or os.path.basename(file_path),
            ocr_text=structured.get("ocr_text") or "",
            summary=structured.get("summary") or "",
            keywords=structured.get("keywords") or [],
            embedding=embedding,
            thumbnail_path=thumb_path if os.path.exists(thumb_path) else None,
            status="done",
            ocr_model=used_model_id,
        )

        # 重算候选链接
        try:
            graph_api.recompute_links_for_note(note_id)
        except Exception as ge:
            logger.warning("链接重算失败 note %s: %s", note_id, ge)

        logger.info("笔记 %s 处理完成 (demo=%s, model=%s)", note_id, is_demo, used_model_id)
        return True

    except Exception as e:
        logger.exception("笔记 %s 处理失败: %s", note_id, e)
        database.update_note_status(note_id, "failed")
        return False
